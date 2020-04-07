import json
import datetime
import csv
import re
import docx
import timeit
from docx.shared import Cm
from docxtpl import DocxTemplate, InlineImage

# Файл с _данными_ состоит из нескольких строк
# с данными по разным машинам. Одна строка - данные по одной машине
datafile_name = "cars.txt"


# Cоздаем словарь из 1(одной) строки передаваемого файла
def get_line_2dict(line_str):
    # # Делаем сплит строки регулярным выражением
    l_dict = dict()
    line_split = re.split(r';\s', line_str.strip('\n'))
    for couple in line_split:
        couple_split = re.split(r':\s', couple)
        l_dict.update({couple_split[0]: couple_split[1]})
    return l_dict


# Делаем переменные из ключей словаря, обьявляем их глобальными
def get_dict_data(l_dict):
    for key, value in l_dict.items():
        globals()[key] = value
    return


# Передаем переменные словаря в качестве аргумента.
# Ключи из файла, также являются именами переменных
# и ключами для вставки в template
def generate_report(**l_dict):
    template = 'template_for_python.docx'
    template = DocxTemplate(template)
    img_size = Cm(17)  # sets the size of the image
    pic = InlineImage(template, Image, img_size)
    l_dict['pic'] = pic  # adds the InlineImage object to the context
    template.render(l_dict)
    global outdocxfile_name
    outdocxfile_name = (Brand + '_' + str(datetime.datetime.now().date()) + '_report.docx')
    template.save(outdocxfile_name)
    return True


# Записываем CSV файл
def save_dict_2csv(l_dict):
    global outcsvfile_name
    outcsvfile_name = (Brand + '_' + str(datetime.datetime.now().date()) + '_report.csv')
    with open(outcsvfile_name, 'w') as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=l_dict.keys())
        writer.writeheader()
        writer.writerow(l_dict)
    return True


# Записываем JSON файл
def save_dict_2json(l_dict):
    global outjsonfile_name
    outjsonfile_name = (Brand + '_' + str(datetime.datetime.now().date()) + '_report.json')
    with open(outjsonfile_name, 'w') as jsonfile:
        json.dump(l_dict, jsonfile)
    return True


# Добавляем в docx документ время затраченное на генерацию отчета
def add_docx_report_time(t_elapsed):
    doc = docx.Document(outdocxfile_name)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    time_str = ('Отчет сгенерирован за ' + str(t_elapsed) + ' сек')
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph(time_str)
    doc.save(outdocxfile_name)
    return


# Добавляем в csv документ время затраченное на генерацию отчета
def add_csv_report_time(t_elapsed):
    # Читаем файл
    with open(outcsvfile_name, 'r') as csvfile:
        l_dict = dict()
        reader = csv.DictReader(csvfile)
        for row in reader:
            l_dict = row
        l_dict.update({'GenTime': str(t_elapsed)})
        # Записываем в файл обновленные данные
        save_dict_2csv(l_dict)
    return


# Добавляем в json документ время затраченное на генерацию отчета
def add_json_report_time(t_elapsed):
    # Читаем файл
    with open(outjsonfile_name, 'r') as jsonfile:
        file_content = jsonfile.read()
        l_dict = json.loads(file_content)
        l_dict.update({'GenTime': str(t_elapsed)})
        # Записываем в файл обновленные данные
        save_dict_2json(l_dict)
    return


# Основное тело
line_dict = dict()
with open(datafile_name, 'r', encoding='utf-8') as textfile:
    for line in textfile:
        line_dict = get_line_2dict(line)
        # Создаем глобальные переменные из созданного словаря
        get_dict_data(line_dict)
        # Генерируем отчет и замеряем время
        time_elapsed = timeit.timeit("generate_report(**line_dict)", setup="from __main__ import generate_report",
                                     number=1, globals=globals())
        # Добавляем время в файл
        add_docx_report_time(time_elapsed)
        # Генерируем отчет и замеряем время
        time_elapsed = timeit.timeit("save_dict_2csv(line_dict)", setup="from __main__ import save_dict_2csv", number=1,
                                     globals=globals())
        # Добавляем время в файл
        add_csv_report_time(time_elapsed)
        # Генерируем отчет и замеряем время
        time_elapsed = (
            timeit.timeit("save_dict_2json(line_dict)", setup="from __main__ import save_dict_2json", number=1,
                          globals=globals()))
        # Добавляем время в файл
        add_json_report_time(time_elapsed)
